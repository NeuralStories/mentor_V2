"""
Parser unificado de documentos para la zona de ingesta.
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Tuple

from core.config import settings

try:
    import fitz
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
except Exception:
    DocxDocument = None
    PackageNotFoundError = Exception

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from PIL import Image, ImageFilter, ImageOps
except Exception:
    Image = None
    ImageFilter = None
    ImageOps = None

logger = logging.getLogger(__name__)

_PAGE_TEXT_MIN_CHARS = 30
_OCR_ZOOM = 2.5


class DocumentParser:
    SUPPORTED_FORMATS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".md": "markdown",
    }

    @staticmethod
    def can_parse(file_path: str) -> bool:
        return Path(file_path).suffix.lower() in DocumentParser.SUPPORTED_FORMATS

    @staticmethod
    def get_format(file_path: str) -> str:
        return DocumentParser.SUPPORTED_FORMATS.get(Path(file_path).suffix.lower(), "unknown")

    @staticmethod
    def validate_file(file_path: str, max_size_mb: int | None = None) -> Tuple[bool, str]:
        try:
            path = Path(file_path)
            if not path.exists():
                return False, "Archivo no encontrado"
            if path.stat().st_size == 0:
                return False, "Archivo vacio"
            if not DocumentParser.can_parse(str(path)):
                return False, f"Formato no soportado: {path.suffix}"

            limit_mb = max_size_mb if max_size_mb is not None else settings.MAX_UPLOAD_SIZE_MB
            size_mb = path.stat().st_size / (1024 * 1024)
            if size_mb > limit_mb:
                return False, f"Archivo demasiado grande ({size_mb:.1f}MB > {limit_mb}MB)"
            return True, "Archivo valido"
        except Exception as exc:
            return False, f"Error validando archivo: {exc}"

    @staticmethod
    def parse_file(file_path: str | Path) -> Tuple[str, Dict[str, Any]]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {path}")
        if not DocumentParser.can_parse(str(path)):
            raise ValueError(f"Formato no soportado: {path.suffix}")

        file_format = DocumentParser.get_format(str(path))
        metadata: Dict[str, Any] = {
            "file_name": path.name,
            "file_path": str(path),
            "file_size": path.stat().st_size,
            "file_format": file_format,
            "upload_date": datetime.utcnow().isoformat(),
            "pages": 1,
            "word_count": 0,
            "char_count": 0,
            "parser": file_format,
            "ocr_used": False,
            "ocr_pages": 0,
        }

        if file_format == "pdf":
            content, extra = DocumentParser._parse_pdf(path)
        elif file_format == "docx":
            content, extra = DocumentParser._parse_docx(path)
        elif file_format == "txt":
            content, extra = DocumentParser._parse_txt(path), {}
        elif file_format == "markdown":
            content, extra = DocumentParser._parse_txt(path), {"parser": "markdown"}
        else:
            raise ValueError(f"Parser no implementado para: {file_format}")

        metadata.update(extra)
        metadata["word_count"] = len(content.split())
        metadata["char_count"] = len(content)
        return content, metadata

    @staticmethod
    def _parse_pdf(path: Path) -> Tuple[str, Dict[str, Any]]:
        if fitz is None:
            raise RuntimeError("PyMuPDF no está instalado")

        parts: list[str] = []
        ocr_pages = 0
        with fitz.open(str(path)) as document:
            pdf_meta = document.metadata or {}
            total_pages = len(document)

            for page in document:
                text = (page.get_text("text") or "").strip()
                if len(text) >= _PAGE_TEXT_MIN_CHARS:
                    parts.append(text)
                    continue

                if settings.OCR_ENABLED and DocumentParser._ocr_available():
                    ocr_text = DocumentParser._ocr_page(page)
                    if ocr_text.strip():
                        parts.append(ocr_text.strip())
                        ocr_pages += 1
                elif text:
                    parts.append(text)

        ocr_used = ocr_pages > 0
        if ocr_used and ocr_pages == total_pages:
            parser_name = "tesseract"
        elif ocr_used:
            parser_name = "pymupdf+ocr"
        else:
            parser_name = "pymupdf"

        return "\n\n".join(part for part in parts if part).strip(), {
            "pages": total_pages,
            "parser": parser_name,
            "ocr_used": ocr_used,
            "ocr_pages": ocr_pages,
            "title": pdf_meta.get("title") or None,
            "author": pdf_meta.get("author") or None,
            "subject": pdf_meta.get("subject") or None,
            "creator": pdf_meta.get("creator") or None,
        }

    @staticmethod
    def _ocr_available() -> bool:
        if pytesseract is None or Image is None:
            logger.warning("OCR habilitado pero faltan dependencias opcionales")
            return False
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
        return True

    @staticmethod
    def _ocr_page(page) -> str:
        try:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(_OCR_ZOOM, _OCR_ZOOM), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            image = DocumentParser._preprocess_for_ocr(image)
            return pytesseract.image_to_string(
                image,
                lang=settings.OCR_LANGUAGE,
                config="--oem 3 --psm 3",
            )
        except Exception as exc:
            logger.warning("OCR falló en página: %s", exc)
            return ""

    @staticmethod
    def _preprocess_for_ocr(image):
        if ImageOps is None:
            return image
        image = ImageOps.grayscale(image)
        image = ImageOps.autocontrast(image)
        if ImageFilter is not None:
            image = image.filter(ImageFilter.MedianFilter(size=3))
        return image

    @staticmethod
    def _parse_docx(path: Path) -> Tuple[str, Dict[str, Any]]:
        if DocxDocument is None:
            raise RuntimeError("python-docx no está instalado")

        try:
            document = DocxDocument(str(path))
        except PackageNotFoundError as exc:
            raise ValueError("Error leyendo DOCX: archivo corrupto o no valido") from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        core = getattr(document, "core_properties", None)
        return "\n".join(parts).strip(), {
            "parser": "docx",
            "pages": 1,
            "title": getattr(core, "title", None) if core else None,
            "author": getattr(core, "author", None) if core else None,
            "subject": getattr(core, "subject", None) if core else None,
        }

    @staticmethod
    def _parse_txt(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"):
            try:
                with open(path, "r", encoding=encoding) as file_handle:
                    return file_handle.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("No se pudo decodificar el archivo de texto")
