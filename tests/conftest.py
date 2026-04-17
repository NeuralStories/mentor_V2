from __future__ import annotations

import io
import os
import shutil
import sqlite3
import time
import uuid
from pathlib import Path
from typing import Iterator

import pytest


def tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def wait_for_terminal_status(client, file_id: str, timeout: float = 45.0, poll: float = 0.25) -> dict:
    deadline = time.time() + timeout
    last = None
    while time.time() < deadline:
        response = client.get(f"/api/knowledge/documents/{file_id}/status")
        if response.status_code == 200:
            last = response.json()
            if last.get("status") in ("ready", "failed"):
                return last
        time.sleep(poll)
    pytest.fail(f"Timeout esperando estado terminal de {file_id}. Último: {last}")


@pytest.fixture(scope="session", autouse=True)
def isolate_data_dir() -> Iterator[Path]:
    root = Path("tests/.tmp/mentor_e2e").resolve()
    shutil.rmtree(root, ignore_errors=True)
    root.mkdir(parents=True, exist_ok=True)
    os.environ["UPLOAD_DIR"] = str(root / "uploads")
    os.environ["INGESTION_DB_PATH"] = str(root / "ingestion.sqlite")
    os.environ["CORS_ORIGINS"] = "http://testserver"
    os.environ["SUPABASE_URL"] = ""
    os.environ["SUPABASE_KEY"] = ""
    os.environ["SUPABASE_SERVICE_KEY"] = ""
    os.environ["SUPABASE_REQUIRED"] = "false"
    yield root


@pytest.fixture
def tmp_path() -> Iterator[Path]:
    path = Path("tests/.tmp/cases") / str(uuid.uuid4())
    path.mkdir(parents=True, exist_ok=True)
    yield path


@pytest.fixture(scope="session")
def client(isolate_data_dir):
    from fastapi.testclient import TestClient
    from core import config as cfg
    from core.ingestion import service as ingestion_service

    cfg.settings = cfg.Settings()
    ingestion_service.get_store.cache_clear()

    from api.main import app

    with TestClient(app) as test_client:
        yield test_client


@pytest.fixture(autouse=True)
def reset_ingestion_store():
    yield
    try:
        from core.config import settings
        with sqlite3.connect(settings.INGESTION_DB_PATH) as connection:
            connection.execute("DELETE FROM ingestion")
    except Exception:
        pass


def _pdf_native_bytes(text: str = "Procedimiento de instalacion de tarima flotante.\nPaso 1: preparar superficie.\nPaso 2: colocar manta aislante.\nPaso 3: montar las lamas.") -> bytes:
    import fitz

    document = fitz.open()
    page = document.new_page()
    y = 72
    for line in text.splitlines():
        page.insert_text((72, y), line, fontsize=12)
        y += 20
    return document.tobytes()


def _pdf_scanned_bytes(text: str = "INSTALACION DE RODAPIE\nMedir, cortar a 45 grados,\nfijar con adhesivo MS polimero.") -> bytes:
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 36)
    except Exception:
        font = ImageFont.load_default()

    y = 120
    for line in text.splitlines():
        draw.text((100, y), line, fill="black", font=font)
        y += 60

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    pdf = fitz.open()
    rect = fitz.Rect(0, 0, image.width, image.height)
    page = pdf.new_page(width=image.width, height=image.height)
    page.insert_image(rect, stream=buffer.getvalue())
    return pdf.tobytes()


def _docx_bytes(text: str = "Manual de seguridad en obra.\nUso obligatorio de EPI.\nProcedimiento de bloqueo y etiquetado.") -> bytes:
    from docx import Document

    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Material"
    table.rows[0].cells[1].text = "Cantidad"
    table.rows[1].cells[0].text = "Tornillos 4x40"
    table.rows[1].cells[1].text = "50"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def pdf_native_bytes() -> bytes:
    return _pdf_native_bytes()


@pytest.fixture
def pdf_scanned_bytes() -> bytes:
    return _pdf_scanned_bytes()


@pytest.fixture
def docx_bytes() -> bytes:
    return _docx_bytes()


@pytest.fixture
def txt_bytes() -> bytes:
    return b"Nota tecnica: comprobar humedad < 3% antes de instalar."


@pytest.fixture
def md_bytes() -> bytes:
    return b"# Checklist\n\n- Humedad < 3%\n- Superficie plana\n- Manta aislante\n"
