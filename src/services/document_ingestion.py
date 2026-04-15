from dataclasses import dataclass
from pathlib import Path
from typing import Optional
import hashlib

from src.core.knowledge_base import KBDocument, DocumentType
from src.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class IngestionResult:
    filename: str
    doc_id: str
    success: bool
    chunks_created: int = 0
    pages_processed: int = 0
    error: str = ""


class DocumentIngestionService:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".xlsx", ".html", ".json"}

    def __init__(self, rag_service=None, knowledge_base=None):
        self.rag = rag_service
        self.kb = knowledge_base

    async def ingest_file(self, filepath: str, metadata: dict = None) -> IngestionResult:
        path = Path(filepath)
        if not path.exists():
            return IngestionResult(filename=path.name, doc_id="", success=False, error=f"File not found: {filepath}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return IngestionResult(filename=path.name, doc_id="", success=False, error=f"Unsupported format: {ext}")

        try:
            content, pages = self._extract_text(path, ext)
            if not content.strip():
                return IngestionResult(filename=path.name, doc_id="", success=False, error="No text content extracted")

            doc_id = hashlib.md5(f"{path.name}{content[:200]}".encode()).hexdigest()[:12]
            metadata = metadata or {}
            title = metadata.get("title", path.stem.replace("_", " ").replace("-", " ").title())
            department = metadata.get("department", "")
            doc_type = metadata.get("type", self._detect_doc_type(path.name, content))
            tags = metadata.get("tags", self._auto_tag(content))

            if self.kb:
                kb_doc = KBDocument(
                    doc_id=doc_id, title=title, content=content,
                    doc_type=DocumentType(doc_type), department=department,
                    tags=tags, version=metadata.get("version", "1.0"),
                    approved_by=metadata.get("approved_by", ""),
                    access_level=metadata.get("access_level", "all"),
                )
                self.kb.add_document(kb_doc)

            chunks_created = 0
            if self.rag:
                chunks_created = self.rag.ingest_document(
                    doc_id=doc_id, title=title, content=content,
                    metadata={"department": department, "doc_type": doc_type, "tags": ",".join(tags), "source_file": path.name},
                )

            logger.info(f"Ingested: {path.name} → {doc_id} ({chunks_created} chunks, {pages} pages)")
            return IngestionResult(filename=path.name, doc_id=doc_id, success=True, chunks_created=chunks_created, pages_processed=pages)

        except Exception as e:
            logger.error(f"Ingestion failed for {path.name}: {e}")
            return IngestionResult(filename=path.name, doc_id="", success=False, error=str(e))

    async def ingest_directory(self, directory: str, recursive: bool = True, default_metadata: dict = None) -> dict:
        dir_path = Path(directory)
        if not dir_path.is_dir():
            return {"error": f"Not a directory: {directory}"}

        pattern = "**/*" if recursive else "*"
        files = [f for f in dir_path.glob(pattern) if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS]

        results = []
        for filepath in files:
            rel_path = filepath.relative_to(dir_path)
            folder_meta = self._metadata_from_path(rel_path)
            meta = {**(default_metadata or {}), **folder_meta}
            result = await self.ingest_file(str(filepath), metadata=meta)
            results.append(result)

        successful = [r for r in results if r.success]
        failed = [r for r in results if not r.success]
        return {
            "total_files": len(files),
            "successful": len(successful),
            "failed": len(failed),
            "total_chunks": sum(r.chunks_created for r in successful),
            "total_pages": sum(r.pages_processed for r in successful),
            "errors": [{"file": r.filename, "error": r.error} for r in failed],
        }

    def _extract_text(self, path: Path, ext: str) -> tuple[str, int]:
        if ext == ".pdf":
            return self._extract_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(path)
        elif ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8"), 1
        elif ext == ".csv":
            return self._extract_csv(path)
        elif ext == ".xlsx":
            return self._extract_xlsx(path)
        elif ext == ".html":
            return self._extract_html(path)
        elif ext == ".json":
            return self._extract_json(path)
        else:
            raise ValueError(f"No extractor for {ext}")

    def _extract_pdf(self, path: Path) -> tuple[str, int]:
        try:
            import pymupdf
            doc = pymupdf.open(str(path))
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            doc.close()
            return "\n\n".join(pages), len(pages)
        except ImportError:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(text)
            return "\n\n".join(pages), len(pages)

    def _extract_docx(self, path: Path) -> tuple[str, int]:
        import docx
        doc = docx.Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))
        return "\n\n".join(paragraphs), 1

    def _extract_csv(self, path: Path) -> tuple[str, int]:
        import csv
        rows = []
        with open(path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows), 1

    def _extract_xlsx(self, path: Path) -> tuple[str, int]:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True)
        sheets_text = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                sheets_text.append(f"Hoja: {sheet}\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets_text), len(sheets_text)

    def _extract_html(self, path: Path) -> tuple[str, int]:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.texts = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style"):
                    self._skip = True

            def handle_endtag(self, tag):
                if tag in ("script", "style"):
                    self._skip = False

            def handle_data(self, data):
                if not self._skip and data.strip():
                    self.texts.append(data.strip())

        html_content = path.read_text(encoding="utf-8")
        extractor = TextExtractor()
        extractor.feed(html_content)
        return "\n".join(extractor.texts), 1

    def _extract_json(self, path: Path) -> tuple[str, int]:
        import json
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, list):
            parts = []
            for item in data:
                if isinstance(item, dict):
                    parts.append("\n".join(f"{k}: {v}" for k, v in item.items()))
                else:
                    parts.append(str(item))
            return "\n\n".join(parts), 1
        elif isinstance(data, dict):
            return "\n".join(f"{k}: {v}" for k, v in data.items()), 1
        return str(data), 1

    def _detect_doc_type(self, filename: str, content: str) -> str:
        name_lower = filename.lower()
        if any(w in name_lower for w in ["politica", "policy", "norma"]):
            return "policy"
        elif any(w in name_lower for w in ["sop", "procedimiento", "proceso", "manual"]):
            return "sop"
        elif any(w in name_lower for w in ["faq", "preguntas"]):
            return "faq"
        elif any(w in name_lower for w in ["guia", "guide", "tutorial"]):
            return "manual"
        return "faq"

    def _auto_tag(self, content: str, max_tags: int = 10) -> list[str]:
        tag_keywords = {
            "vacaciones": ["vacaciones", "días libres", "descanso"],
            "nómina": ["nómina", "salario", "sueldo", "pago"],
            "seguridad": ["seguridad", "emergencia", "evacuación"],
            "sistemas": ["sistema", "software", "acceso", "contraseña"],
            "reembolso": ["reembolso", "gasto", "factura", "viáticos"],
            "capacitación": ["capacitación", "curso", "entrenamiento"],
            "beneficios": ["beneficio", "seguro", "prestación"],
        }
        content_lower = content.lower()
        found_tags = []
        for tag, keywords in tag_keywords.items():
            if any(kw in content_lower for kw in keywords):
                found_tags.append(tag)
        return found_tags[:max_tags]

    def _metadata_from_path(self, rel_path: Path) -> dict:
        parts = rel_path.parts
        meta = {}
        if len(parts) >= 2:
            meta["department"] = parts[0]
        return meta
